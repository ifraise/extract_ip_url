#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从 .doc .docx .xls .xlsx 提取
  IP      -> ip.txt
  URL     -> url.txt
  domain  -> domain.txt
usage: python extract_target.py
"""
import re
import os
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
try:
    import xlrd  # 1.2.0
except ImportError:
    xlrd = None

IPv4_RE   = re.compile(r'\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b')

# 匹配 URL
URL_RE    = re.compile(
    r'https?://'
    r'(?:[a-zA-Z0-9-]+\.)+[a-zA-Z0-9-]+'
    r'(?::\d+)?'
    r'(?:/[a-zA-Z0-9._~:/?#[\]@!$&\'()*+,;=-]*)?'
)

# 提取域名
DOMAIN_RE = re.compile(
    r'(?<=^|(?<=[^a-zA-Z0-9-]))'
    r'(?:[a-zA-Z0-9-]+\.)+'
    r'[a-zA-Z]{2,}'
    r'(?=$|(?=[^a-zA-Z0-9-]))'
)

def _run(cmd, timeout=30):
    try:
        ret = subprocess.run(cmd, shell=True, capture_output=True,
                             text=True, timeout=timeout)
        if ret.returncode == 0:
            return ret.stdout
    except Exception:
        pass
    return None

def libreoffice_convert(folder: Path, dst: Path):
    if not shutil.which('soffice'):
        return []
    cmd = f'soffice --headless --convert-to docx,xlsx --outdir "{dst}" "{folder}"/*'
    _run(cmd, timeout=60)
    return list(dst.glob('*.docx')) + list(dst.glob('*.xlsx'))

def extract_docx(path: Path):
    doc = Document(path)
    text = []
    for p in doc.paragraphs:
        text.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for c in row.cells:
                text.append(c.text)
    return '\n'.join(text)

def extract_doc_antiword(path: Path):
    return _run(f'antiword "{path}"') or ''

def extract_xlsx(path: Path):
    wb = load_workbook(path, read_only=True, data_only=True)
    text = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for v in row:
                if v is not None:
                    text.append(str(v))
    return '\n'.join(text)

def extract_xls_xlrd(path: Path):
    if xlrd is None:
        return ''
    wb = xlrd.open_workbook(path)
    text = []
    for ws in wb.sheets():
        for r in range(ws.nrows):
            for c in range(ws.ncols):
                v = ws.cell_value(r, c)
                if v:
                    text.append(str(v))
    return '\n'.join(text)

def main():
    root = Path.cwd()
    ip_pool, url_pool, domain_pool = set(), set(), set()

    todo = list(root.rglob('*.doc')) + list(root.rglob('*.docx')) + \
           list(root.rglob('*.xls')) + list(root.rglob('*.xlsx'))
    if not todo:
        print('[-] 未找到任何 Office 文件！(灬ꈍ ꈍ灬)')
        return

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        old_docs = [f for f in todo if f.suffix == '.doc']
        old_xls  = [f for f in todo if f.suffix == '.xls']
        if old_docs or old_xls:
            print('[*] LibreOffice 转换旧格式…')
            for batch in (old_docs, old_xls):
                if not batch:
                    continue
                d = tmp / ('doc' if batch == old_docs else 'xls')
                d.mkdir()
                for f in batch:
                    shutil.copy(f, d)
                out = d / 'out'
                libreoffice_convert(d, out)
                todo.extend(out.glob('*.docx'))
                todo.extend(out.glob('*.xlsx'))

        for f in todo:
            print(f'[+] 处理 {f}')
            try:
                ext = f.suffix.lower()
                if ext == '.docx':
                    txt = extract_docx(f)
                elif ext == '.xlsx':
                    txt = extract_xlsx(f)
                elif ext == '.doc':
                    txt = extract_doc_antiword(f)
                elif ext == '.xls':
                    txt = extract_xls_xlrd(f)
                else:
                    continue
                ip_pool.update(IPv4_RE.findall(txt))
                url_pool.update(URL_RE.findall(txt))
                for d in DOMAIN_RE.findall(txt):
                    domain_pool.add(d)
                print(f'    IP {len(ip_pool)} 条，URL {len(url_pool)} 条，域名 {len(domain_pool)} 条')
            except Exception as e:
                print(f'    [!] 失败: {e}')

    with open('ip.txt', 'w', encoding='utf-8') as f:
        for ip in sorted(ip_pool):
            f.write(ip + '\n')
    with open('url.txt', 'w', encoding='utf-8') as f:
        for url in sorted(url_pool):
            f.write(url + '\n')
    with open('domain.txt', 'w', encoding='utf-8') as f:
        for d in sorted(domain_pool):
            f.write(d + '\n')

    print('\n[*] 全部完成！(*￣︶￣)')
    print(f'    IP  : {len(ip_pool)} 条 -> ip.txt')
    print(f'    URL : {len(url_pool)} 条 -> url.txt')
    print(f'    域名: {len(domain_pool)} 条 -> domain.txt')

if __name__ == '__main__':
    main()
