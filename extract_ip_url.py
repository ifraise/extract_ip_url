import re
import pandas as pd
import argparse
import os
from docx import Document

# 改进的正则表达式
URL_REGEX = r"https?://[^\s<>\"'{}()]+"
IP_REGEX = r"\b(?:[0-9]{1,3}\.){3}[0-9]{1,3}\b"

def extract_ip_and_url(text):
    """
    从文本中提取 IP 和 URL，使用正则表达式匹配。
    """
    ips = re.findall(IP_REGEX, text)
    urls = re.findall(URL_REGEX, text)
    
    # 过滤不合法的 IP 地址
    ips = [ip for ip in ips if all(0 <= int(part) <= 255 for part in ip.split('.'))]
    return ips, urls

def process_excel(file_path):
    """
    处理 Excel 文件，提取 IP 和 URL。
    """
    ips = set()
    urls = set()
    df = pd.read_excel(file_path, header=None)
    for _, row in df.iterrows():
        for cell in row:
            extracted_ips, extracted_urls = extract_ip_and_url(str(cell))
            ips.update(extracted_ips)
            urls.update(extracted_urls)
    return list(ips), list(urls)

def process_csv(file_path):
    """
    处理 CSV 文件，提取 IP 和 URL。
    """
    ips = set()
    urls = set()
    df = pd.read_csv(file_path, header=None)
    for _, row in df.iterrows():
        for cell in row:
            extracted_ips, extracted_urls = extract_ip_and_url(str(cell))
            ips.update(extracted_ips)
            urls.update(extracted_urls)
    return list(ips), list(urls)

def process_txt(file_path):
    """
    处理 TXT 文件，提取 IP 和 URL。
    """
    ips = set()
    urls = set()
    with open(file_path, 'r', encoding='utf-8') as f:
        for line in f:
            extracted_ips, extracted_urls = extract_ip_and_url(line)
            ips.update(extracted_ips)
            urls.update(extracted_urls)
    return list(ips), list(urls)

def process_word(file_path):
    """
    处理 Word 文件，提取 IP 和 URL。
    """
    ips = set()
    urls = set()
    doc = Document(file_path)
    
    # 提取段落中的内容
    for para in doc.paragraphs:
        extracted_ips, extracted_urls = extract_ip_and_url(para.text)
        ips.update(extracted_ips)
        urls.update(extracted_urls)

    # 提取表格中的内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extracted_ips, extracted_urls = extract_ip_and_url(cell.text)
                ips.update(extracted_ips)
                urls.update(extracted_urls)
    
    return list(ips), list(urls)

def save_results(ips, urls, output_file):
    """
    保存提取结果到 Excel 文件。
    """
    max_len = max(len(ips), len(urls))
    ips.extend([None] * (max_len - len(ips)))
    urls.extend([None] * (max_len - len(urls)))
    df = pd.DataFrame({'IP': ips, 'URL': urls})
    df.to_excel(output_file, index=False)
    print(f"结果已保存到 {output_file}")

def main():
    parser = argparse.ArgumentParser(description="提取文件中的 IP 和 URL")
    parser.add_argument('-r', '--read', required=True, help="输入文件路径（支持 TXT、CSV、Excel 和 Word 文件）")
    parser.add_argument('-o', '--output', help="输出文件路径（默认生成 Excel 文件）")
    args = parser.parse_args()

    file_path = args.read

    # 根据文件扩展名选择处理方法
    if file_path.endswith('.txt'):
        print(f"正在处理 TXT 文件：{file_path}")
        ips, urls = process_txt(file_path)
    elif file_path.endswith('.xlsx') or file_path.endswith('.xls'):
        print(f"正在处理 Excel 文件：{file_path}")
        ips, urls = process_excel(file_path)
    elif file_path.endswith('.csv'):
        print(f"正在处理 CSV 文件：{file_path}")
        ips, urls = process_csv(file_path)
    elif file_path.endswith('.docx'):
        print(f"正在处理 Word 文件：{file_path}")
        ips, urls = process_word(file_path)
    else:
        print("不支持的文件格式，仅支持 TXT、CSV、Excel 和 Word 文件。")
        return

    # 设置输出文件名
    output_file = args.output if args.output else f"{os.path.splitext(file_path)[0]}_result.xlsx"
    save_results(ips, urls, output_file)

if __name__ == '__main__':
    main()
